import json
import shutil
import ssl
import sys
from pathlib import Path
from urllib import request

import certifi
import filetype
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Cm, Pt
from PIL import Image
from PySide6 import QtGui, QtWidgets
from PySide6.QtCore import Qt, Signal
from PySide6.QtWidgets import QMessageBox

import resources_rc

from .AboutUI import Ui_AboutWindow
from .CheckUI import Ui_CheckWindows
from .CustomWidgets import LoadImage, QListWidgetItemPlus
from .MainUI import Ui_MainWindow
from .ManualUI import Ui_ManualWindow


class MainWindow_controller(QtWidgets.QMainWindow):
    send_update_image_count = Signal()

    def __init__(self):
        super().__init__()

        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.version = "2.2"

        self.show_usage_warrning_window()

        self.set_statusbar()
        self.setup_control()

    def show_usage_warrning_window(self):
        """
        show the usage_warrning message
        """
        messagebox = QMessageBox(self)
        messagebox.setWindowTitle("警告")
        messagebox.setText("任何圖片在進入此程式使用前，請務必做好備份，並且使用副本來操作，以避免一切會造成圖片原檔消失的可能性。")
        messagebox.setIcon(QMessageBox.Icon.Critical)
        messagebox.addButton("我了解了", QMessageBox.ButtonRole.AcceptRole)
        messagebox.addButton("我還沒備份", QMessageBox.ButtonRole.RejectRole)

        result = messagebox.exec()
        if result == 1:
            import sys

            sys.exit()

    def set_statusbar(self):
        """
        initialize the statusbar
        """
        self.img_count = QtWidgets.QLabel()
        self.img_count.setText(f"列表中共有 0 張圖片")
        self.status = QtWidgets.QLabel()
        self.status.setText("讀取狀態： 無")
        self.ui.statusbar.addPermanentWidget(self.img_count)
        self.ui.statusbar.addPermanentWidget(self.status)

    def setup_control(self):
        """
        initialize global variable and connect slot functions
        """
        self.is_thumbnail_mode = self.ui.thumnail_rtn.isChecked()

        self.progress_vaule = 0

        self.send_update_image_count.connect(self.update_image_count)

        self.ui.menu_about_manual.triggered.connect(self.show_manual_window)
        self.ui.menu_about_update.triggered.connect(self.check_update)
        self.ui.menu_about_about.triggered.connect(self.show_about_window)

        self.ui.list_rtn.toggled.connect(self.change_image_mode)
        self.ui.thumnail_rtn.toggled.connect(self.change_image_mode)

        self.ui.path_btn.clicked.connect(self.add_folder_img)
        self.ui.items_list.send_new_item.connect(self.add_drag_img)

        self.ui.clear_list_btn.clicked.connect(self.clear_list)
        self.ui.items_list.itemDoubleClicked.connect(self.show_big_img_window)

        self.ui.start_btn.clicked.connect(self.show_setting_window)
        self.ui.exit_btn.clicked.connect(self.app_exit)

    def update_image_count(self):
        """
        updata image count in the statusbar
        """
        self.img_count.setText(f"列表中共有 {self.ui.items_list.count()} 張圖片")

    def show_manual_window(self):
        """
        open manual window
        """
        self.ManualWindows = QtWidgets.QWidget()
        self.manual_ui = Ui_ManualWindow()
        self.manual_ui.setupUi(self.ManualWindows)
        self.ManualWindows.setFixedSize(800, 600)
        self.manual_ui.exit_btn.clicked.connect(lambda: self.ManualWindows.close())
        self.ManualWindows.show()

    def check_update(self):
        github_release_url: str = (
            "https://github.com/scbmark/auto_image_report/releases/latest"
        )
        github_release_url_api: str = (
            "https://api.github.com/repos/scbmark/auto_image_report/releases/latest"
        )

        req = request.Request(github_release_url_api)

        try:
            context = ssl.create_default_context(cafile=certifi.where())
            with request.urlopen(req, context=context) as response:
                res = json.load(response)
                latest_version = res["tag_name"]
        except:
            messagebox = QMessageBox(self)
            messagebox.warning(self, "更新", f"網路連線失敗")
            return

        if latest_version != self.version:
            messagebox = QMessageBox(self)
            messagebox.setWindowTitle("更新")
            messagebox.setText(f"發現更新\n目前版本： {self.version}\n最新版本： {res['tag_name']}")
            messagebox.setIcon(QMessageBox.Icon.Information)
            messagebox.addButton("現在更新", QMessageBox.ButtonRole.AcceptRole)
            messagebox.addButton("不要更新", QMessageBox.ButtonRole.RejectRole)

            result = messagebox.exec()
            if result == 0:
                import webbrowser

                webbrowser.open(github_release_url)
            else:
                pass
        else:
            messagebox = QMessageBox(self)
            messagebox.information(self, "檢查更新", f"目前為最新版本\n目前版本： {self.version}")

    def show_about_window(self):
        """
        open about window
        """
        self.AboutWindows = QtWidgets.QWidget()
        self.about_ui = Ui_AboutWindow()
        self.about_ui.setupUi(self.AboutWindows)
        self.AboutWindows.setFixedSize(400, 600)
        self.about_ui.exit_btn.clicked.connect(lambda: self.AboutWindows.close())
        self.AboutWindows.show()

    def change_image_mode(self):
        """
        chage the image mode between thumbnail and text
        """
        self.is_thumbnail_mode = self.ui.thumnail_rtn.isChecked()

    def add_folder_img(self):
        """
        add images in folder
        """
        dir_path = self.open_img_dir()

        if dir_path == "":
            return

        pictures, invalid_files = self.load_image(dir_path)

        if invalid_files:
            invalid_files = "\n".join(invalid_files)
            messabebox = QtWidgets.QMessageBox(self)
            messabebox.warning(self, "Error", f"以下檔案不支援\n {invalid_files}")

        if not pictures:
            messabebox = QMessageBox(self)
            messabebox.warning(self, "Error", "無偵測到圖片")
            return

        if self.is_thumbnail_mode is None:
            self.show_is_thumbnail_window()

        if self.is_thumbnail_mode:
            self.load_img_thread = LoadImage(pictures)
            self.load_img_thread.send_img_item.connect(self.add_img_item_to_list)
            self.load_img_thread.send_progress.connect(self.update_progress)

            self.load_img_thread.start()
        else:
            for count, pic in enumerate(pictures, start=1):
                self.update_progress(f"{count} / {len(pictures)}")
                self.add_text_item_to_list(pic)

            self.update_progress("無")

    def open_img_dir(self) -> str:
        """
        show the file dialog and get the folder's path
        """
        dir_path = QtWidgets.QFileDialog.getExistingDirectory(caption="選取資料夾")
        self.ui.path_box.setText(dir_path)

        return dir_path

    def load_image(self, dir_path: str) -> tuple[list[str], list[str]]:
        """
        load image's filename from the folder's path and return the sorted filename list and the invalid filename list
        """
        all_file: list[str] = list()
        pictures: list[str] = list()
        invalid_links: list[str] = list()

        support_format = ["jpg", "png", "bmp", "tif"]

        path = Path(dir_path)
        all_file.extend(path.glob("*"))

        for file in all_file:
            if Path(file).is_dir():
                continue

            file_format = filetype.guess(file)
            file_name = str(file)
            if file_format is None:
                invalid_links.append(file_name)
            elif file_format.extension in support_format:
                pictures.append(file_name)
            else:
                invalid_links.append(file_name)

        return sorted(pictures), invalid_links

    def add_drag_img(self, pictures: list):
        """
        add images from drag and drop
        """
        if self.is_thumbnail_mode is None:
            self.show_is_thumbnail_window()

        if self.is_thumbnail_mode:
            self.load_img_thread = LoadImage(pictures)
            self.load_img_thread.send_img_item.connect(self.add_img_item_to_list)
            self.load_img_thread.send_progress.connect(self.update_progress)

            self.load_img_thread.start()
        else:
            for count, pic in enumerate(pictures, start=1):
                self.update_progress(f"{count} / {len(pictures)}")
                self.add_text_item_to_list(pic)

            self.update_progress("無")

    def show_is_thumbnail_window(self):
        """
        show the dialog to select the image mode
        """
        messagebox = QMessageBox()
        messagebox.setWindowTitle("選擇模式")
        messagebox.setText("選擇是否載入縮圖，以方便調整圖片順序")
        messagebox.setIcon(QMessageBox.Icon.Warning)
        messagebox.addButton("載入", QMessageBox.ButtonRole.YesRole)
        messagebox.addButton("不載入", QMessageBox.ButtonRole.NoRole)

        result = messagebox.exec()

        if result == 0:
            self.is_thumbnail_mode = True

        else:
            self.is_thumbnail_mode = False

    def update_progress(self, progress: str):
        """
        updata image loadin progress in the statusbar
        """
        self.status.setText(f"讀取狀態： {progress}")

    def add_img_item_to_list(self, item_tuple: tuple):
        """
        create thumbnail mode QListItemWidget and add it to the QListWiget
        """
        img_w = item_tuple[0].width()
        img_h = item_tuple[0].height()

        MAXIMUM_HEIGHT = 64
        ratio = MAXIMUM_HEIGHT / img_h

        item = QListWidgetItemPlus(item_tuple[1])
        item_widget = QtWidgets.QWidget()

        pic_name_lb = QtWidgets.QLabel()
        pic_name_lb.setText(Path(item_tuple[1]).name)
        item_layout = QtWidgets.QHBoxLayout()

        pic_view = QtWidgets.QGraphicsView(self)
        pic_view.setGeometry(0, 0, int(img_w * ratio), MAXIMUM_HEIGHT)
        pic_sence = QtWidgets.QGraphicsScene()
        pic_sence.setSceneRect(0, 0, int(img_w * ratio), MAXIMUM_HEIGHT)

        pic_sence.addPixmap(item_tuple[0])
        pic_view.setScene(pic_sence)

        delete_button = QtWidgets.QPushButton("")
        delete_button.setIcon(QtGui.QIcon(QtGui.QPixmap(":/statics/close.svg")))
        delete_button.clicked.connect(lambda: self.delete_item(item))

        item_layout.addWidget(pic_name_lb)
        item_layout.addWidget(pic_view)
        item_layout.addStretch()
        item_layout.addWidget(delete_button)

        item_widget.setLayout(item_layout)
        item.setSizeHint(item_widget.sizeHint())

        self.ui.items_list.addItem(item)
        self.ui.items_list.setItemWidget(item, item_widget)
        self.ui.items_list.repaint()
        self.send_update_image_count.emit()

    def add_text_item_to_list(self, path: str):
        """
        create text mode QListItemWidget and add it to the QListWiget
        """
        item = QListWidgetItemPlus(path)
        item_widget = QtWidgets.QWidget()

        pic_name_lb = QtWidgets.QLabel()
        pic_name_lb.setText(Path(path).name)
        item_layout = QtWidgets.QHBoxLayout()

        delete_button = QtWidgets.QPushButton("")
        delete_button.setIcon(QtGui.QIcon(QtGui.QPixmap(":/statics/close.svg")))
        delete_button.clicked.connect(lambda: self.delete_item(item))

        item_layout.addWidget(pic_name_lb)
        item_layout.addStretch()
        item_layout.addWidget(delete_button)

        item_widget.setLayout(item_layout)
        item.setSizeHint(item_widget.sizeHint())

        self.ui.items_list.addItem(item)
        self.ui.items_list.setItemWidget(item, item_widget)
        self.ui.items_list.repaint()
        self.send_update_image_count.emit()

    def delete_item(self, item: QListWidgetItemPlus):
        """
        get the QListWidgetItem and remove it from QListWidget
        """
        index = self.ui.items_list.row(item)
        self.ui.items_list.takeItem(index)
        self.send_update_image_count.emit()

    def clear_list(self):
        """
        remove all QListWidgetItems in the QListWidget
        """
        self.ui.items_list.clear()
        self.send_update_image_count.emit()

    def show_big_img_window(self, item: QListWidgetItemPlus):
        """
        open and show the big image
        """
        MAXIMUM_WIDTH = 768
        MAXIMUM_HEIGHT = 512

        with Image.open(item.text) as img:
            img.thumbnail(size=(MAXIMUM_WIDTH, MAXIMUM_HEIGHT))

            img_raw_data = img.convert("RGB").tobytes("raw", "RGB")
            img = QtGui.QImage(
                img_raw_data,
                img.size[0],
                img.size[1],
                img.size[0] * 3,
                QtGui.QImage.Format.Format_RGB888,
            )

        img = QtGui.QPixmap.fromImage(img)

        self.img_widget = QtWidgets.QLabel()

        self.img_widget.setWindowFlags(
            Qt.WindowType.Window | Qt.WindowType.WindowStaysOnTopHint
        )

        self.img_widget.setAlignment(
            Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignTop
        )

        self.img_widget.setGeometry(0, 0, int(img.width()), int(img.height()))
        self.img_widget.setWindowTitle(f"{Path(item.text).name}")
        self.img_widget.setPixmap(img)
        self.img_widget.show()

    def show_setting_window(self):
        """
        show the report configuration window
        """
        picture_counts = self.ui.items_list.count()
        picture_list = list()

        for index in range(picture_counts):
            item_obj = self.ui.items_list.item(index)
            picture_list.append(item_obj.text)

        if not picture_list:
            messabebox = QMessageBox(self)
            messabebox.warning(self, "Error", "列表中無圖片")
            return

        self.CheckWindows = QtWidgets.QWidget()
        self.check_ui = Ui_CheckWindows()
        self.check_ui.setupUi(self.CheckWindows)
        self.CheckWindows.setFixedSize(410, 510)

        self.check_ui.path_btn.clicked.connect(self.select_output_dir)
        self.check_ui.custom_num_rtn.toggled.connect(self.toggle_custom_num)
        self.check_ui.start_generate_btn.clicked.connect(
            lambda: self.start_generate(picture_list)
        )

        self.CheckWindows.show()

    def select_output_dir(self):
        """
        show the file dialog and get the path where the report is exported
        """
        dir_path = QtWidgets.QFileDialog.getExistingDirectory(caption="選取資料夾")
        self.check_ui.path_box.setText(dir_path)

    def toggle_custom_num(self):
        """
        enable or disable the optional numbering features by custom_num_rtn's state
        """
        if self.check_ui.custom_num_rtn.isChecked():
            self.check_ui.prefix_box.setEnabled(True)
            self.check_ui.subfix_box.setEnabled(True)
            self.check_ui.rule_box.setEnabled(True)
        else:
            self.check_ui.prefix_box.setEnabled(False)
            self.check_ui.subfix_box.setEnabled(False)
            self.check_ui.rule_box.setEnabled(False)

    def start_generate(self, picture_list: list):
        """
        the main function to generate the report
        """
        self.progress_value = 0
        configs = self.get_configs()

        if not (configs["path"] and configs["filename"] and configs["title"]):
            messabebox = QMessageBox(self)
            messabebox.warning(self, "Error", "參數不完整，請填寫參數")
            return

        adjust_numbers = list()
        if configs["custom_num"]["number_rule"] == [""]:
            adjust_numbers = [str(i) for i in range(1, len(picture_list) + 1)]

        else:
            for num in configs["custom_num"]["number_rule"]:
                if "-" in num:
                    start = int(num.split("-")[0])
                    end = int(num.split("-")[1])
                    adjust_numbers.extend([str(i) for i in range(start, end + 1)])
                else:
                    adjust_numbers.append(num)

        if len(picture_list) != len(adjust_numbers):
            messabebox = QMessageBox(self)
            messabebox.warning(self, "Error", "自訂編號規則和圖片數量不一致")
            return

        if configs["compress"]:
            picture_list = self.img_compression(picture_list, configs["path"])

        self.check_ui.path_btn.setDisabled(True)
        self.check_ui.start_generate_btn.setDisabled(True)

        try:
            self.generate_docx(
                path=configs["path"],
                filename=configs["filename"],
                title=configs["title"],
                custom_num=configs["custom_num"],
                pictures=picture_list,
                adjust_numbers=adjust_numbers,
            )

            messabebox = QMessageBox(self)
            messabebox.information(
                self, "Success", f"報告匯出成功\n插入了 {len(picture_list)} 張圖片"
            )

            compress_path = Path().joinpath(configs["path"], "compressed")
            if compress_path.exists():
                shutil.rmtree(compress_path)

            self.CheckWindows.close()
        except:
            compress_path = Path().joinpath(configs["path"], "compressed")
            if compress_path.exists():
                shutil.rmtree(compress_path)

            messabebox = QMessageBox(self)
            messabebox.warning(self, "Error", "報告生成失敗，請聯繫開發者")

            self.CheckWindows.close()

    def get_configs(self) -> dict:
        """
        return all configuration in dict format
        """
        configs = dict()

        configs["path"] = self.check_ui.path_box.text()
        configs["filename"] = self.check_ui.filename_box.text()
        configs["title"] = self.check_ui.title_box.text()
        configs["number_prefix"] = self.check_ui.prefix_box.text()
        configs["number_subfix"] = self.check_ui.subfix_box.text()

        if self.check_ui.custom_num_rtn.isChecked():
            configs["custom_num"] = {
                "is_custom_num": True,
                "number_prefix": self.check_ui.prefix_box.text(),
                "number_subfix": self.check_ui.subfix_box.text(),
            }
            if self.check_ui.rule_box.text() == None:
                configs["custom_num"]["number_rule"] = [""]
            else:
                configs["custom_num"]["number_rule"] = (
                    self.check_ui.rule_box.text().strip().split(",")
                )
        else:
            configs["custom_num"] = {
                "is_custom_num": False,
                "number_prefix": "",
                "number_subfix": "",
                "number_rule": [""],
            }

        configs["compress"] = self.check_ui.is_compress.isChecked()

        return configs

    def img_compression(self, filename_list: list, root_path: str) -> list[str]:
        """
        load and compress the image in the filename list
        """
        compress_path = Path().joinpath(root_path, "compressed")

        if not compress_path.exists():
            compress_path.mkdir()
        else:
            shutil.rmtree(compress_path)
            compress_path.mkdir()

        self.progress_value = 0
        self.check_ui.progressBar.setRange(0, len(filename_list))

        new_filename_list: list[str] = list()
        for index, file in enumerate(filename_list, start=1):
            self.check_ui.current_progress_lb.setText(
                f"圖片壓縮中... , {index}/{len(filename_list)}"
            )
            self.check_ui.current_progress_lb.repaint()

            base_name = Path(file).name
            compressed_img_filename = str(
                Path().joinpath(compress_path, f"c-{base_name}")
            )

            with Image.open(file) as img:
                img.save(compressed_img_filename, format="JPEG", quality=80)

            new_filename_list.append(compressed_img_filename)

            self.progress_value += 1
            self.check_ui.progressBar.setValue(self.progress_value)

        self.progress_value = 0
        return new_filename_list

    def generate_docx(
        self,
        path: str,
        filename: str,
        title: str,
        custom_num: dict,
        pictures: list,
        adjust_numbers: list,
    ):
        """
        use python-docx package to insert image and generate report
        """
        doc = Document()

        section = doc.sections[0]

        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.add_run(title).add_break()

        for run in paragraph.runs:
            run.bold = True

        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in paragraph.runs:
            run.font.name = "標楷體"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
            run.font.size = Pt(20)

        self.progress_value = 0
        self.check_ui.progressBar.setRange(0, len(pictures))

        for index, num in enumerate(adjust_numbers):
            pic_para = doc.add_paragraph()
            pic_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_run = pic_para.add_run("")

            pic = pic_run.add_picture(pictures[index], width=Cm(13.5))

            MAXIMUM_HEIGHT = 9.7
            ratio = MAXIMUM_HEIGHT / pic.height.cm

            pic.width = Cm(pic.width.cm * ratio)
            pic.height = Cm(MAXIMUM_HEIGHT)

            self.check_ui.current_progress_lb.setText(
                f"圖片插入中... , {index+1}/{len(pictures)}"
            )
            print(f"插入第{index+1}張圖片，共有{len(pictures)}張")

            if custom_num["is_custom_num"]:
                num_para = doc.add_paragraph()
                num_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                num_para.add_run(
                    f"照片{custom_num['number_prefix']+num+custom_num['number_subfix']}"
                ).add_break()

                for run in num_para.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
                    run.font.size = Pt(14)

            else:
                num_para = doc.add_paragraph()
                num_para._element.get_or_add_pPr()

                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), "autonumber")

                num_para._element.pPr.append(pStyle)

                num_para.add_run("").add_break()

                for run in num_para.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
                    run.font.size = Pt(14)

            self.progress_value += 1
            self.check_ui.progressBar.setValue(self.progress_value)

        print("插入完成，正在輸出報告")

        self.check_ui.current_progress_lb.setText("檔案正在匯出...")
        self.check_ui.progress_lb.repaint()

        doc_path = str(Path().joinpath(path, f"{filename}.docx"))
        doc.save(doc_path)

        self.check_ui.current_progress_lb.setText("匯出完成")

    def app_exit(self):
        """
        close this program
        """
        sys.exit()
